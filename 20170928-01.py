# -*- coding: utf-8 -*-
"""
python-docx lib test

@author: Bryson Xue

@Note: 
    1. 測試如何產生docx file
    2. 測試如何做簡單的格式控制
    3. pip install python-docx
"""

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

document = Document()

#文件格式風格設定
obj_styles = document.styles
obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(10)
#obj_font.name = 'Times New Roman'
obj_font.name = '標楷體'

#文件title
document.add_heading('Heading Title', level=0)

#文件內文段落
paragraph = document.add_paragraph()

#中文字似乎套用字型都無效，英文則無問題
run = paragraph.add_run('test test test', style = 'CommentsStyle')
#run = paragraph.add_run(u'中文字中文字中文字', style = 'CommentsStyle')
run.italic = True	#斜體
run.bold = True		#粗體

#加入表格
t = document.add_table(rows=4, cols=2)

#可用table styles 參考
#http://python-docx.readthedocs.io/en/latest/user/styles-understanding.html
#t.style = 'Light List'
#t.style = 'Light Grid'
t.style = 'Table Grid'	

t.cell(0, 0).text = u'title A'
t.cell(0, 1).text = u'title B'
t.cell(1, 0).text = u'unit 1'
t.cell(1, 1).text = u'unit 2'
t.cell(2, 0).text = u'unit 3'
t.cell(2, 1).text = u'unit 4'
t.cell(3, 0).text = u''
t.cell(3, 1).text = u'unit 5'

#儲存格合併
a = t.cell(2, 0)
b = t.cell(3, 0)
a.merge(b)

#文件內文段落2
paragraph2 = document.add_paragraph()

run = paragraph2.add_run('aaaaaaaaaaaaaaaaaaaaaa')
run.add_break()	#斷行
run = paragraph2.add_run('bbbbbbbbbbbbbbbbbbbbbb')
document.add_page_break()	#分頁

#文件內文段落3
paragraph3 = document.add_paragraph()
run = paragraph3.add_run('cccccccccccccccccccccc')
run.italic = True
document.add_page_break()


#文件內文段落4
paragraph3 = document.add_paragraph()
run = paragraph3.add_run('dddddddddddddddddddddd')
run.bold = True
document.add_page_break()

#產生檔案
document.save('py_doc.docx')