# -*- coding: utf-8 -*-
"""
讀取EXCEL資料，輸出成WORD表格

@author: Bryson Xue

@Note: 

@Ref:
	https://stackoverflow.com/questions/20297332/python-pandas-dataframe-retrieve-number-of-columns
	https://stackoverflow.com/questions/40596518/writing-a-python-pandas-dataframe-to-word-document

"""
import openpyxl as px
import pandas as pd
from docx import Document

#讀取資料來源EXCEL FILE
W = px.load_workbook('test_data.xlsx')
p = W.get_sheet_by_name(name = '工作表1')

rows = p.rows
columns = p.columns

content = []
for row in rows:
    line = [col.value for col in row]
    content.append(line)

#print(content)
df = pd.DataFrame(content[1:], columns = content[0])
#print(df)

#取的df資料筆數
row_cnt = df.shape[0]
#print(str(df.shape[0]))

#取的df資料欄位數
col_cnt = df.shape[1]
#print(str(df.shape[1]))

document = Document()

#文件內文段落
paragraph = document.add_paragraph()

#加入表格
t = document.add_table(rows = row_cnt+1, cols = col_cnt)
t.style = 'Light Grid'

# add the header rows.
for j in range(df.shape[-1]):
    t.cell(0,j).text = df.columns[j]

# add the rest of the data frame
for i in range(df.shape[0]):
    for j in range(df.shape[-1]):
    	t.cell(i+1,j).text = str(df.values[i,j])

# save the doc
document.save('./aaa.docx')